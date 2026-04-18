# -*- coding: utf-8 -*-
"""
RTBIO Statement of Work v3.0 DOCX Generator
Generates a professional Korean SOW document using python-docx.
Run with: python -X utf8 build_v3_docx.py
"""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# ── Constants ──────────────────────────────────────────────────────────────────
OUTPUT_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                           "RTBIO_과업지시서_v3.0.docx")
FONT_NAME = "맑은 고딕"
PRIMARY = RGBColor(0x1B, 0x3A, 0x5C)    # #1B3A5C
ACCENT  = RGBColor(0x00, 0xA8, 0xB5)    # #00A8B5
WHITE   = RGBColor(0xFF, 0xFF, 0xFF)
BLACK   = RGBColor(0x00, 0x00, 0x00)
ALT_ROW = "F2F6FA"
HDR_BG  = "1B3A5C"
ACCENT_HEX = "00A8B5"
LIGHT_GRAY = "E8E8E8"


# ── Helper functions ──────────────────────────────────────────────────────────

def set_cell_shading(cell, color_hex):
    """Set background shading for a table cell."""
    shading = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_margins(cell, top=60, bottom=60, left=100, right=100):
    """Set cell margins in twips."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'  <w:top w:w="{top}" w:type="dxa"/>'
        f'  <w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'  <w:start w:w="{left}" w:type="dxa"/>'
        f'  <w:end w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcMar_old = tcPr.find(qn('w:tcMar'))
    if tcMar_old is not None:
        tcPr.remove(tcMar_old)
    tcPr.append(tcMar)


def set_cell_vertical_alignment(cell, align="center"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    vAlign = parse_xml(
        f'<w:vAlign {nsdecls("w")} w:val="{align}"/>'
    )
    old = tcPr.find(qn('w:vAlign'))
    if old is not None:
        tcPr.remove(old)
    tcPr.append(vAlign)


def set_run_font(run, size=10, bold=False, color=BLACK, font_name=FONT_NAME):
    """Configure a run's font properties."""
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.name = font_name
    r = run._element
    rPr = r.find(qn('w:rPr'))
    if rPr is None:
        rPr = parse_xml(f'<w:rPr {nsdecls("w")}/>')
        r.insert(0, rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)


def add_paragraph(doc, text, size=10, bold=False, color=BLACK, alignment=None,
                  space_before=0, space_after=6):
    """Add a paragraph with specified formatting."""
    p = doc.add_paragraph()
    if alignment is not None:
        p.alignment = alignment
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    return p


def add_section_header(doc, number, title, level=1):
    """Add a numbered section header."""
    if level == 1:
        p = add_paragraph(doc, f"{number}. {title}", size=24, bold=True,
                          color=PRIMARY, space_before=24, space_after=12)
        # Add bottom border
        pPr = p._element.get_or_add_pPr()
        pBdr = parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'  <w:bottom w:val="single" w:sz="8" w:space="4" w:color="{HDR_BG}"/>'
            f'</w:pBdr>'
        )
        pPr.append(pBdr)
    elif level == 2:
        p = add_paragraph(doc, f"{number} {title}", size=18, bold=True,
                          color=PRIMARY, space_before=18, space_after=8)
    return p


def create_table(doc, headers, rows, col_widths=None):
    """Create a formatted table with header row and alternating colors."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Set column widths if provided
    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)

    # Header row
    for i, header_text in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header_text)
        set_run_font(run, size=9, bold=True, color=WHITE)
        set_cell_shading(cell, HDR_BG)
        set_cell_margins(cell)
        set_cell_vertical_alignment(cell)

    # Data rows
    for row_idx, row_data in enumerate(rows):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(cell_text))
            set_run_font(run, size=9, bold=False, color=BLACK)
            if row_idx % 2 == 1:
                set_cell_shading(cell, ALT_ROW)
            set_cell_margins(cell)
            set_cell_vertical_alignment(cell)

    # Remove extra spacing in all cell paragraphs
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(2)
                paragraph.paragraph_format.space_after = Pt(2)

    return table


def add_callout_box(doc, text):
    """Add a styled callout/note box."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(12)
    # Add left border and background via XML
    pPr = p._element.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:left w:val="single" w:sz="24" w:space="8" w:color="{ACCENT_HEX}"/>'
        f'  <w:top w:val="single" w:sz="4" w:space="2" w:color="{LIGHT_GRAY}"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="2" w:color="{LIGHT_GRAY}"/>'
        f'  <w:right w:val="single" w:sz="4" w:space="2" w:color="{LIGHT_GRAY}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F0F8FF" w:val="clear"/>')
    pPr.append(shd)
    # Indentation
    ind = parse_xml(f'<w:ind {nsdecls("w")} w:left="200" w:right="200"/>')
    pPr.append(ind)

    icon_run = p.add_run("NOTE: ")
    set_run_font(icon_run, size=9, bold=True, color=ACCENT)
    text_run = p.add_run(text)
    set_run_font(text_run, size=9, bold=False, color=RGBColor(0x33, 0x33, 0x33))


def add_page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    br_elem = parse_xml(f'<w:br {nsdecls("w")} w:type="page"/>')
    run._element.append(br_elem)


# ── Main document builder ────────────────────────────────────────────────────

def build_document():
    doc = Document()

    # ── Page setup ────────────────────────────────────────────────────────
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # Set default font
    style = doc.styles['Normal']
    style.font.name = FONT_NAME
    style.font.size = Pt(10)
    style.font.color.rgb = BLACK
    rPr = style.element.find(qn('w:rPr'))
    if rPr is None:
        rPr = parse_xml(f'<w:rPr {nsdecls("w")}/>')
        style.element.append(rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), FONT_NAME)

    # ══════════════════════════════════════════════════════════════════════
    # COVER / HEADER
    # ══════════════════════════════════════════════════════════════════════

    # Spacer
    for _ in range(4):
        add_paragraph(doc, "", size=10, space_after=0)

    # Top accent line
    p_line = doc.add_paragraph()
    p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr = p_line._element.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="24" w:space="1" w:color="{ACCENT_HEX}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)

    # Title
    add_paragraph(doc, "", size=10, space_after=0)
    p_title = add_paragraph(doc, "", size=36, bold=True, color=PRIMARY,
                            alignment=WD_ALIGN_PARAGRAPH.CENTER,
                            space_before=20, space_after=6)
    run_title = p_title.add_run("\uACFC\uC5C5\uC9C0\uC2DC\uC11C")  # 과업지시서
    set_run_font(run_title, size=36, bold=True, color=PRIMARY)

    # Subtitle
    p_sub = add_paragraph(doc, "", size=18, color=ACCENT,
                          alignment=WD_ALIGN_PARAGRAPH.CENTER,
                          space_before=6, space_after=12)
    run_sub = p_sub.add_run("RTBIO \uC5C5\uBB34 \uC790\uB3D9\uD654 \uC2DC\uC2A4\uD15C")  # RTBIO 업무 자동화 시스템
    set_run_font(run_sub, size=18, bold=False, color=ACCENT)

    # Bottom accent line
    p_line2 = doc.add_paragraph()
    p_line2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr2 = p_line2._element.get_or_add_pPr()
    pBdr2 = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="24" w:space="1" w:color="{ACCENT_HEX}"/>'
        f'</w:pBdr>'
    )
    pPr2.append(pBdr2)

    # Spacer
    add_paragraph(doc, "", size=10, space_after=0)

    # Version & Date
    add_paragraph(doc, "Version 3.0", size=14, bold=True, color=PRIMARY,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=12, space_after=4)
    add_paragraph(doc, "2026.04.10", size=12, color=RGBColor(0x66, 0x66, 0x66),
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=6)

    # Spacer
    for _ in range(6):
        add_paragraph(doc, "", size=10, space_after=0)

    # Company info on cover
    add_paragraph(doc, "\uD558\uB9C8\uB2E4\uB7A9\uC2A4 | HAMADA LABS", size=12, bold=True, color=PRIMARY,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=4)
    add_paragraph(doc, "victoria@hamadalabs.com", size=10,
                  color=RGBColor(0x66, 0x66, 0x66),
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=0)

    # ── Page break ────────────────────────────────────────────────────────
    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════════════
    # 1. 문서 정보
    # ══════════════════════════════════════════════════════════════════════
    add_section_header(doc, "1", "\uBB38\uC11C \uC815\uBCF4")

    # Two-column info table
    info_table = doc.add_table(rows=4, cols=4)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    info_table.style = 'Table Grid'

    left_labels  = ["\uACF5\uAE09\uC790", "\uB300\uD45C", "\uC5F0\uB77D\uCC98", "\uC774\uBA54\uC77C"]
    left_values  = ["\uD558\uB9C8\uB2E4\uB7A9\uC2A4", "\uBC29\uC2B9\uC560", "010-8888-0180", "victoria@hamadalabs.com"]
    right_labels = ["\uC218\uC2E0", "\uACAC\uC801\uC77C", "\uC720\uD6A8\uAE30\uAC04", ""]
    right_values = ["\uC54C\uD2F0\uBC14\uC774\uC624", "2026\uB144 4\uC6D4 10\uC77C", "\uBC1C\uD589\uC77C\uB85C\uBD80\uD130 30\uC77C", ""]

    for i in range(4):
        # Left label
        cell_ll = info_table.rows[i].cells[0]
        cell_ll.text = ""
        p = cell_ll.paragraphs[0]
        run = p.add_run(left_labels[i])
        set_run_font(run, size=9, bold=True, color=PRIMARY)
        set_cell_shading(cell_ll, "E8EEF5")
        set_cell_margins(cell_ll)

        # Left value
        cell_lv = info_table.rows[i].cells[1]
        cell_lv.text = ""
        p = cell_lv.paragraphs[0]
        run = p.add_run(left_values[i])
        set_run_font(run, size=9, bold=False, color=BLACK)
        set_cell_margins(cell_lv)

        # Right label
        cell_rl = info_table.rows[i].cells[2]
        cell_rl.text = ""
        p = cell_rl.paragraphs[0]
        run = p.add_run(right_labels[i])
        set_run_font(run, size=9, bold=True, color=PRIMARY)
        if right_labels[i]:
            set_cell_shading(cell_rl, "E8EEF5")
        set_cell_margins(cell_rl)

        # Right value
        cell_rv = info_table.rows[i].cells[3]
        cell_rv.text = ""
        p = cell_rv.paragraphs[0]
        run = p.add_run(right_values[i])
        set_run_font(run, size=9, bold=False, color=BLACK)
        set_cell_margins(cell_rv)

    # Set widths for info table
    for row in info_table.rows:
        row.cells[0].width = Cm(2.8)
        row.cells[1].width = Cm(5.2)
        row.cells[2].width = Cm(2.8)
        row.cells[3].width = Cm(5.2)
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(2)
                paragraph.paragraph_format.space_after = Pt(2)

    add_paragraph(doc, "", size=6, space_after=0)

    # ══════════════════════════════════════════════════════════════════════
    # 2. 프로젝트 개요
    # ══════════════════════════════════════════════════════════════════════
    add_section_header(doc, "2", "\uD504\uB85C\uC81D\uD2B8 \uAC1C\uC694")

    overview_items = [
        ("AS-IS", "\uCE74\uCE74\uC624\uD1A1 \uBC1C\uC8FC \uC811\uC218 \u2192 \uC5BC\uB9C8\uC5D0\uC694 \uC218\uAE30 \uC785\uB825 \u2192 \uC5D1\uC140 \uB9C8\uAC10 \uCC98\uB9AC"),
        ("TO-BE", "\uC6F9 \uAE30\uBC18 \uD1B5\uD569 \uC2DC\uC2A4\uD15C\uC73C\uB85C \uC804\uD658. \uBC1C\uC8FC \uC811\uC218~\uC815\uC0B0\uAE4C\uC9C0 \uC790\uB3D9\uD654."),
        ("\uBAA9\uD45C", "\uC5C5\uBB34 \uD6A8\uC728\uC131 50% \uC774\uC0C1 \uD5A5\uC0C1, \uC5BC\uB9C8\uC5D0\uC694 \uC644\uC804 \uB300\uCCB4"),
        ("\uD50C\uB7AB\uD3FC", "\uC6F9 \uAE30\uBC18 (\uC571 X), \uBAA8\uBC14\uC77C \uBC18\uC751\uD615, \uD074\uB77C\uC6B0\uB4DC(AWS)"),
    ]

    for label, desc in overview_items:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        # Bullet-like prefix
        run_label = p.add_run(f"  {label}:  ")
        set_run_font(run_label, size=10, bold=True, color=PRIMARY)
        run_desc = p.add_run(desc)
        set_run_font(run_desc, size=10, bold=False, color=BLACK)

    add_paragraph(doc, "", size=6, space_after=0)

    # ══════════════════════════════════════════════════════════════════════
    # 3. 사용자 정의
    # ══════════════════════════════════════════════════════════════════════
    add_section_header(doc, "3", "\uC0AC\uC6A9\uC790 \uC815\uC758")

    user_headers = ["\uC5ED\uD560", "\uC0AC\uC6A9\uC790", "\uC8FC\uC694 \uC5C5\uBB34", "\uC811\uADFC \uBC94\uC704"]
    user_rows = [
        ["\uAC70\uB798\uCC98 (\uC678\uBD80)", "\uB300\uB9AC\uC810/\uBCD1\uC6D0 \uB2F4\uB2F9\uC790 (40\uB300+ \uB2E4\uC218)", "\uBC1C\uC8FC \uC785\uB825, \uB0B4\uC5ED \uC870\uD68C, \uCD9C\uACE0 \uD655\uC778", "\uBCF8\uC778 \uBC1C\uC8FC/\uB2E8\uAC00\uB9CC"],
        ["\uACBD\uC601\uC9C0\uC6D0\uD300", "\uB0B4\uBD80 1~2\uBA85", "\uBC1C\uC8FC \uD655\uC778, \uAC70\uB798\uBA85\uC138\uC11C, \uB9C8\uAC10 \uC6D0\uC7A5, \uBCF4\uACE0\uC11C", "\uC804\uCCB4 \uAC70\uB798\uCC98 \uB9E4\uCD9C/\uBBF8\uC218\uAE08"],
        ["\uD488\uC9C8\uAD00\uB9AC\uD300", "\uB0B4\uBD80 6\uBA85", "\uBC1C\uC8FC \uD655\uC815, \uB2F4\uB2F9\uC790 \uBC30\uC815, \uD3EC\uC7A5/\uCD9C\uACE0, \uC7AC\uACE0", "\uBC1C\uC8FC \uB0B4\uC5ED, \uC7AC\uACE0, \uCD9C\uACE0 \uC0C1\uD0DC"],
        ["\uC601\uC5C5\uD300", "\uB0B4\uBD80 N\uBA85", "\uC0AC\uC6A9\uB7C9 \uC785\uB825, \uAC70\uB798\uCC98 \uBC29\uBB38 \uC2E4\uC801", "\uB2F4\uB2F9 \uAC70\uB798\uCC98 \uC0AC\uC6A9\uB7C9/\uB9E4\uCD9C \uC870\uD68C"],
        ["\uB300\uD45C/\uC784\uC6D0", "\uB300\uD45C, \uC804\uBB34", "\uB300\uC2DC\uBCF4\uB4DC, \uC804\uCCB4 \uB9E4\uCD9C/\uC7AC\uACE0 \uD604\uD669", "\uC804\uCCB4 \uC2DC\uC2A4\uD15C \uC77D\uAE30 (\uCD5C\uC0C1\uC704)"],
    ]
    create_table(doc, user_headers, user_rows, col_widths=[3.0, 4.0, 5.0, 4.0])

    add_paragraph(doc, "", size=6, space_after=0)

    # ══════════════════════════════════════════════════════════════════════
    # 4. 제공 범위 — 기능 상세 명세
    # ══════════════════════════════════════════════════════════════════════
    add_section_header(doc, "4", "\uC81C\uACF5 \uBC94\uC704 \u2014 \uAE30\uB2A5 \uC0C1\uC138 \uBA85\uC138")

    add_callout_box(doc,
        "\uBCF8 \uAE30\uB2A5 \uBAA9\uB85D\uC740 1\uCC28 \uBBF8\uD305(2026.04.10) \uAE30\uC900 \uC815\uB9AC \uB0B4\uC6A9\uC774\uBA70, "
        "\uC138\uBD80 \uD654\uBA74 \uAD6C\uC131 \uBC0F UX\uB294 2\uCC28 \uBBF8\uD305 \uD6C4 \uCD5C\uC885 \uD655\uC815\uD569\uB2C8\uB2E4."
    )

    # ── 4.1 ──────────────────────────────────────────────────────────────
    add_section_header(doc, "4.1", "1\uCC28 \u2014 \uACBD\uC601\uC9C0\uC6D0\uD300 \uC790\uB3D9\uD654 (\uBA54\uC778 \uC2DC\uC2A4\uD15C + DB \uC778\uD504\uB77C)", level=2)

    r1_headers = ["ID", "\uAE30\uB2A5\uBA85", "\uC0C1\uC138 \uC124\uBA85", "\uBE44\uACE0"]
    r1_rows = [
        ["R01", "\uAC70\uB798\uCC98 \uBC1C\uC8FC\uD3FC",
         "\uAE30\uC874 \uC5D1\uC140 \uC591\uC2DD \uAE30\uBC18, \uC804 \uC81C\uD488 \uB178\uCD9C + \uC218\uB7C9 \uC785\uB825 \uBC29\uC2DD. "
         "\uB85C\uADF8\uC778 \uC2DC \uAC70\uB798\uCC98\uBCC4 \uD560\uC778 \uB2E8\uAC00 \uC801\uC6A9\uB41C \uC81C\uD488 \uBAA9\uB85D \uD45C\uC2DC. "
         "\uC81C\uD488\uAD70\uBCC4(\uBB34\uB98E/\uC0C1\uC9C0/\uD558\uC9C0/\uC2A4\uD504\uB9B0\uD2B8 \uB4F1) \uD0ED \uAD6C\uBD84. "
         "\uC218\uB7C9 \uC785\uB825 \uD6C4 \uD569\uACC4 \uC790\uB3D9 \uACC4\uC0B0.",
         "\uD654\uBA74 \uAD6C\uC131 2\uCC28 \uBBF8\uD305 \uD6C4 \uD655\uC815"],
        ["R02", "\uAC70\uB798\uCC98 \uC778\uC99D",
         "RTBIO\uAC00 \uC0DD\uC131\uD55C ID(\uB2F4\uB2F9\uC790\uBA85)/PW(\uAD00\uB9AC\uBC88\uD638)\uB85C \uB85C\uADF8\uC778. "
         "\uB85C\uADF8\uC778 \uC2DC \uD560\uC778\uC728, \uBC30\uC1A1 \uC8FC\uC18C, \uACB0\uC81C \uBC29\uC2DD \uC790\uB3D9 \uB85C\uB4DC. "
         "\uB2E4\uB978 \uAC70\uB798\uCC98 \uC815\uBCF4 \uC811\uADFC \uBD88\uAC00.", ""],
        ["R03", "\uBC1C\uC8FC \u2192 DB \uC790\uB3D9 \uBC18\uC601",
         "\uC6F9 \uBC1C\uC8FC \uC811\uC218 \uC989\uC2DC DB \uC800\uC7A5, \uC218\uAE30 \uC785\uB825 \uC644\uC804 \uC81C\uAC70. "
         "\uBC1C\uC8FC\uBC88\uD638 \uC790\uB3D9 \uC0DD\uC131. "
         "\uAC70\uB798\uCC98/\uC81C\uD488/\uC218\uB7C9/\uB2E8\uAC00/\uD560\uC778\uC728/\uD569\uACC4 \uC790\uB3D9 \uACC4\uC0B0.", ""],
        ["R04", "\uBC1C\uC8FC \uD655\uC815 (\uB2F9\uC77C/\uC775\uC77C)",
         "\uD488\uC9C8\uAD00\uB9AC\uD300\uC774 \uBC1C\uC8FC \uD655\uC778 \u2192 \uB2F9\uC77C\uCD9C\uACE0/\uC775\uC77C\uCD9C\uACE0 \uC120\uD0DD. "
         "\uD655\uC815 \uC2DC \uAC70\uB798\uCC98 \uC6F9\uC5D0 \uC0C1\uD0DC \uBC18\uC601. "
         "15:30 \uC774\uD6C4 \uC811\uC218\uBD84\uC740 \uC775\uC77C \uC790\uB3D9 \uBD84\uB958.", ""],
        ["R05", "\uAC70\uB798\uCC98\uBCC4 \uD560\uC778\uC728 \uAD00\uB9AC",
         "\uC81C\uD488\uAD70 \u00D7 \uAC70\uB798\uCC98\uBCC4 \uD560\uC778\uC728 \uC124\uC815. "
         "\uACE0\uC815\uB2E8\uAC00 > \uC81C\uD488\uAD70 \uD560\uC778\uC728 > \uAE30\uBCF8\uB2E8\uAC00 \uC6B0\uC120\uC21C\uC704. "
         "\uBC1C\uC8FC\uD3FC/\uAC70\uB798\uBA85\uC138\uC11C\uC5D0 \uC790\uB3D9 \uC801\uC6A9.",
         "\uD560\uC778\uC728 \uAD6C\uC870 RTBIO \uC790\uB8CC \uC218\uB839 \uD6C4 \uD655\uC815"],
        ["R06", "\uAC70\uB798\uBA85\uC138\uC11C \uC790\uB3D9 \uC0DD\uC131/\uBC1C\uC1A1",
         "\uCD9C\uACE0 \uC644\uB8CC \uC2DC PDF \uC790\uB3D9 \uC0DD\uC131 + \uC774\uBA54\uC77C \uBC1C\uC1A1. "
         "\uAC70\uB798\uCC98 \uC591\uC2DD/RTBIO \uC591\uC2DD \uAD6C\uBD84(\uAC70\uB798\uCC98 \uC124\uC815). "
         "\uBE44\uACE0\uB780\uC5D0 \uBC30\uC1A1\uBC29\uBC95, \uC9C1\uC1A1 \uC2DC \uBC30\uC1A1\uC9C0\uBA85 \uC790\uB3D9 \uAE30\uC785.", ""],
        ["R07", "\uC0C1\uD488 & \uC7AC\uACE0 \uAD00\uB9AC",
         "\uC81C\uD488 \uB4F1\uB85D/\uC218\uC815(\uC81C\uD488\uAD70, \uC0AC\uC774\uC988, \uC21F/\uB871, \uD3B8\uCE21/\uC591\uCE21). "
         "\uCD9C\uACE0 \uC2DC \uC790\uB3D9 \uC7AC\uACE0 \uCC28\uAC10. "
         "\uC785\uACE0 \uB4F1\uB85D \uC2DC \uC7AC\uACE0 \uAC00\uC0B0. "
         "\uC7AC\uACE0 \uBD80\uC871 \uC2DC \uAD00\uB9AC\uC790 \uC54C\uB9BC.",
         "\uC138\uD2B8 \uC0C1\uD488 \uACC4\uC0B0 \uB85C\uC9C1 \uD655\uC778 \uD544\uC694"],
        ["R08", "\uCD9C\uACE0 \uC0C1\uD0DC \uAD00\uB9AC",
         "\uB2F4\uB2F9\uC790 \uBC30\uC815(6\uBA85) + \uC791\uC5C5 \uB2E8\uACC4\uBCC4 \uC0C1\uD0DC: "
         "\uB300\uAE30 \u2192 \uBC14\uCF54\uB4DC \uCD9C\uB825 \u2192 \uC138\uD305 \u2192 \uD3EC\uC7A5 \u2192 \uC1A1\uC7A5 \u2192 \uCD9C\uACE0 \uC644\uB8CC. "
         "\uB204\uAC00/\uC5B8\uC81C \uCC98\uB9AC\uD588\uB294\uC9C0 \uC774\uB825 \uAE30\uB85D.", ""],
        ["R09", "\uBD80\uC11C\uBCC4 \uAD8C\uD55C \uAD00\uB9AC",
         "\uACBD\uC601\uC9C0\uC6D0: \uC804\uCCB4 \uB9E4\uCD9C/\uC815\uC0B0. "
         "\uD488\uC9C8\uAD00\uB9AC: \uBC1C\uC8FC/\uC7AC\uACE0/\uCD9C\uACE0. "
         "\uC601\uC5C5: \uC0AC\uC6A9\uB7C9 \uC785\uB825, \uB2F4\uB2F9 \uAC70\uB798\uCC98. "
         "\uB300\uD45C: \uC804\uCCB4 \uB300\uC2DC\uBCF4\uB4DC.", ""],
        ["R10", "\uBC30\uC1A1\uC9C0 \uAD00\uB9AC",
         "\uAC70\uB798\uCC98\uBCC4 \uAE30\uBCF8 \uBC30\uC1A1 \uC8FC\uC18C \uC800\uC7A5. "
         "\uBC1C\uC8FC \uC2DC \"\uB2E4\uB978 \uC8FC\uC18C\" \uC120\uD0DD \u2192 \uC9C1\uC811 \uC785\uB825(\uBCD1\uC6D0 \uC9C1\uC1A1). "
         "\uBC30\uC1A1 \uBC29\uBC95: \uD0DD\uBC30/\uBC29\uBB38 \uC218\uB839/\uD035.", ""],
    ]
    create_table(doc, r1_headers, r1_rows, col_widths=[1.2, 3.2, 8.6, 3.0])

    add_paragraph(doc, "", size=6, space_after=0)

    # ── 4.2 ──────────────────────────────────────────────────────────────
    add_section_header(doc, "4.2", "2\uCC28 \u2014 \uD488\uC9C8\uAD00\uB9AC\uD300 \uC7AC\uACE0 \uC790\uB3D9\uD654", level=2)

    r2_headers = ["ID", "\uAE30\uB2A5\uBA85", "\uC0C1\uC138 \uC124\uBA85", "\uBE44\uACE0"]
    r2_rows = [
        ["R11", "\uBC1C\uC8FC \uC218\uC815 (\uC2DC\uAC04 \uC81C\uD55C)",
         "\uD0DD\uBC30 \uB9C8\uAC10(15:30) \uC804\uAE4C\uC9C0 \uAC70\uB798\uCC98\uAC00 \uC9C1\uC811 \uC218\uB7C9 \uBCC0\uACBD \uAC00\uB2A5. "
         "\uB9C8\uAC10 \uD6C4 \uC218\uC815 \uBD88\uAC00 \uC7A0\uAE08. "
         "\uCD9C\uACE0 \"\uD3EC\uC7A5\" \uC774\uD6C4 \uC218\uC815 \uBD88\uAC00.", ""],
        ["R15", "\uBB34\uB8CC\uBC30\uC1A1 \uC548\uB0B4",
         "\uBC1C\uC8FC\uD3FC\uC5D0\uC11C \uC218\uB7C9 \uC785\uB825 \uC2DC \uBB34\uB8CC\uBC30\uC1A1 \uC5EC\uBD80 \uD45C\uC2DC + "
         "\"N\uAC1C \uB354 \uCD94\uAC00\uD558\uBA74 \uBB34\uB8CC\uBC30\uC1A1\" \uC54C\uB9BC. "
         "\uC81C\uD488\uAD70\uBCC4 \uC870\uAC74 \uC124\uC815.",
         "\uBB34\uB8CC\uBC30\uC1A1 \uAE30\uC900 RTBIO \uC790\uB8CC \uC218\uB839 \uD6C4 \uD655\uC815"],
        ["R16", "\uC77C\uC77C \uC7AC\uACE0 \uD604\uD669 \uC790\uB3D9 \uC9D1\uACC4",
         "\uC81C\uD488\uAD70\uBCC4 \uBC1C\uC8FC\uB7C9/\uCD9C\uACE0\uB7C9/\uC794\uC5EC \uC7AC\uACE0 \uC790\uB3D9 \uC815\uB9AC. "
         "\uD1F4\uADFC \uC2DC \uC5D1\uC140 \uC815\uB9AC \uB300\uCCB4. "
         "\uC81C\uD488\uAD70\uBCC4 \uCEA1\uCC98 \uACF5\uC720 \uB300\uCCB4.",
         "\uB9AC\uD3EC\uD2B8 \uC591\uC2DD 2\uCC28 \uBBF8\uD305 \uD6C4 \uD655\uC815"],
        ["R18", "\uC0D8\uD50C/\uBBF8\uD305\uC6A9 \uCD9C\uACE0",
         "\uBCC4\uB3C4 \uCD9C\uACE0 \uC720\uD615 \"\uC0D8\uD50C\". "
         "\uB9E4\uCD9C \uBBF8\uBC18\uC601 + \uC7AC\uACE0\uC5D0\uC11C\uB9CC \uCC28\uAC10. "
         "\uC6D0\uAC00 \uAE30\uB85D \uB0A8\uAE40.", ""],
    ]
    create_table(doc, r2_headers, r2_rows, col_widths=[1.2, 3.2, 8.6, 3.0])

    add_paragraph(doc, "", size=6, space_after=0)

    # ── 4.3 ──────────────────────────────────────────────────────────────
    add_section_header(doc, "4.3", "3\uCC28 \u2014 \uC601\uC5C5\uBD80 \uB300\uC2DC\uBCF4\uB4DC", level=2)

    r3_headers = ["ID", "\uAE30\uB2A5\uBA85", "\uC0C1\uC138 \uC124\uBA85", "\uBE44\uACE0"]
    r3_rows = [
        ["R12", "\uB9C8\uAC10 \uC6D0\uC7A5 \uC790\uB3D9 \uC0DD\uC131",
         "\uAC70\uB798\uCC98\uBCC4 \uB9C8\uAC10 \uAE30\uAC04(\uB300\uB9AC\uC810 1~25\uC77C, \uBCD1\uC6D0 \uC804\uB2EC 26~\uB2F9\uC6D4 25) \uC124\uC815 \u2192 "
         "\uC6D0\uC7A5 \uC790\uB3D9 \uACC4\uC0B0 + \uC794\uC561 \uC815\uB9AC + \uBA54\uC77C \uBC1C\uC1A1.",
         "\uB9C8\uAC10 \uC6D0\uC7A5 \uC591\uC2DD RTBIO \uC790\uB8CC \uC218\uB839 \uD6C4 \uD655\uC815"],
        ["R13", "\uAC70\uB798\uCC98\uBCC4 \uACB0\uC81C \uBC29\uC2DD \uAD00\uB9AC",
         "\uB2F9\uC6D4 \uB9D0 \uCE74\uB4DC, \uC0AC\uC6A9\uB7C9 \uCE74\uB4DC, \uACC4\uC88C\uC774\uCCB4, N\uAC1C\uC6D4 \uD6C4 \uACB0\uC81C. "
         "\uACB0\uC81C \uBC29\uC2DD\uBCC4 \uACC4\uC0B0\uC11C \uBC1C\uD589 \uC5EC\uBD80 \uC790\uB3D9 \uD310\uB2E8.", ""],
        ["R14", "\uC77C\uC77C \uBCF4\uACE0\uC11C \uC790\uB3D9 \uC0DD\uC131",
         "\uC5C5\uCCB4\uBCC4 \uBC1C\uC8FC \uC218\uB7C9/\uAE08\uC561 + \uB2F9\uC77C \uCD1D \uB9E4\uCD9C \uC790\uB3D9 \uC9D1\uACC4. "
         "\uB300\uD45C/\uC784\uC6D0 \uB300\uC2DC\uBCF4\uB4DC\uC5D0 \uD45C\uC2DC.",
         "\uBCF4\uACE0\uC11C \uC591\uC2DD RTBIO \uC790\uB8CC \uC218\uB839 \uD6C4 \uD655\uC815"],
        ["R17", "\uC601\uC5C5\uD300 \uC0AC\uC6A9\uB7C9 \uC785\uB825",
         "\uBCD1\uC6D0 \uBC29\uBB38 \uD6C4 \uC2E4\uC81C \uC0AC\uC6A9\uB7C9 \uC6F9\uC5D0 \uC785\uB825 \u2192 \uB9C8\uAC10 \uC6D0\uC7A5\uC5D0 \uC790\uB3D9 \uBC18\uC601. "
         "\uC0AC\uC6A9\uB7C9 \uBBF8\uC785\uB825 \uC2DC \uD574\uB2F9 \uBCD1\uC6D0 \uB9C8\uAC10 \uBCF4\uB958.", ""],
        ["R22", "\uAE30\uAC04\uBCC4 \uB9E4\uCD9C \uBCF4\uACE0\uC11C",
         "\uC5C5\uCCB4\uBCC4 \uAE30\uAC04 \uC790\uC720 \uC124\uC815. "
         "\uC81C\uD488\uAD70\uBCC4 \uB9E4\uCD9C/\uBC1C\uC8FC\uB7C9 \uD1B5\uACC4 \uB9AC\uD3EC\uD2B8.",
         "2\uCC28 \uBBF8\uD305 \uD6C4 \uC0C1\uC138 \uD655\uC815"],
    ]
    create_table(doc, r3_headers, r3_rows, col_widths=[1.2, 3.2, 8.6, 3.0])

    add_paragraph(doc, "", size=6, space_after=0)

    # ── 4.4 ──────────────────────────────────────────────────────────────
    add_section_header(doc, "4.4", "\uCD94\uD6C4 \uAC80\uD1A0 (2\uCC28 \uBBF8\uD305 \uD6C4 \uACB0\uC815)", level=2)

    r4_headers = ["ID", "\uAE30\uB2A5\uBA85", "\uC0C1\uC138 \uC124\uBA85", "\uC0C1\uD0DC"]
    r4_rows = [
        ["R19", "\uCE74\uCE74\uC624\uD1A1 \uC54C\uB9BC",
         "\uBC1C\uC8FC \uD655\uC815/\uCD9C\uACE0 \uC644\uB8CC \uC2DC \uAC70\uB798\uCC98\uC5D0 \uCE74\uD1A1 \uC54C\uB9BC.",
         "\uBE44\uC988\uB2C8\uC2A4 \uACC4\uC815 \uBE44\uC6A9 \uD655\uC778 \uD6C4 \uACB0\uC815"],
        ["R20", "\uC7AC\uACE0 \uC218\uC694 \uC608\uCE21 \uC54C\uB9BC",
         "\uACFC\uAC70 \uBC1C\uC8FC \uD328\uD134 \uBD84\uC11D \u2192 \uC7AC\uACE0 \uBD80\uC871 \uC704\uD5D8 \uC0AC\uC804 \uC54C\uB9BC.",
         "AI \uAE30\uBC18, 1\uCC28 \uB0A9\uD488 \uD6C4 \uB370\uC774\uD130 \uCD95\uC801 \uD544\uC694"],
        ["R21", "\uAC70\uB798\uCC98 \uC6F9 \uC6D0\uC7A5 \uC870\uD68C",
         "\uAC70\uB798\uCC98\uAC00 \uC9C1\uC811 \uC6F9\uC5D0\uC11C \uC6D0\uC7A5/\uBBF8\uC218\uAE08 \uC870\uD68C.",
         "2\uCC28 \uBBF8\uD305 \uD6C4 \uACB0\uC815"],
        ["R23", "\uD488\uC808 \uC2DC \uC790\uB3D9 \uBC1C\uC8FC \uCC28\uB2E8",
         "\uC7AC\uACE0 0\uC77C \uB54C \uBC1C\uC8FC \uCC28\uB2E8. "
         "\uB300\uD45C\uB2D8 \uC758\uACAC \uBBF8\uD655\uC815 (\uC77C\uB2E8 \uBC1B\uACE0 \uD6C4\uCC98\uB9AC \uAC00\uB2A5\uC131).",
         "RTBIO \uB300\uD45C \uD655\uC778 \uD544\uC694"],
    ]
    create_table(doc, r4_headers, r4_rows, col_widths=[1.2, 3.2, 8.6, 3.0])

    add_paragraph(doc, "", size=6, space_after=0)

    # ══════════════════════════════════════════════════════════════════════
    # 5. 비즈니스 룰
    # ══════════════════════════════════════════════════════════════════════
    add_page_break(doc)
    add_section_header(doc, "5", "\uBE44\uC988\uB2C8\uC2A4 \uB8F0")

    # ── 발주 규칙 ──
    add_paragraph(doc, "\uBC1C\uC8FC \uADDC\uCE59", size=14, bold=True, color=PRIMARY,
                  space_before=12, space_after=6)

    br_order_headers = ["BR", "\uADDC\uCE59\uBA85", "\uC0C1\uC138"]
    br_order_rows = [
        ["01", "\uD0DD\uBC30 \uB9C8\uAC10 \uC2DC\uAC04",
         "\uC624\uD6C4 15:30 \uC774\uD6C4 \uC811\uC218 \uBC1C\uC8FC\uB294 \uC775\uC77C \uCD9C\uACE0. \uB9C8\uAC10 \uC2DC\uAC04\uC740 \uAD00\uB9AC\uC790 \uC124\uC815 \uAC00\uB2A5."],
        ["02", "\uBC1C\uC8FC \uC218\uC815 \uC81C\uD55C",
         "\uCD9C\uACE0 \uC0C1\uD0DC \"\uB300\uAE30\" \uB610\uB294 \"\uBC14\uCF54\uB4DC \uCD9C\uB825\" \uB2E8\uACC4\uC77C \uB54C\uB9CC \uC218\uC815 \uAC00\uB2A5. \"\uD3EC\uC7A5\" \uC774\uD6C4 \uC218\uC815 \uBD88\uAC00."],
        ["03", "\uD560\uC778\uC728 \uC801\uC6A9 \uC6B0\uC120\uC21C\uC704",
         "\uAC70\uB798\uCC98\uBCC4 \uACE0\uC815\uB2E8\uAC00 > \uAC70\uB798\uCC98\uBCC4 \uC81C\uD488\uAD70 \uD560\uC778\uC728 > \uAE30\uBCF8 \uB2E8\uAC00."],
        ["04", "\uBB34\uB8CC\uBC30\uC1A1 \uAE30\uC900",
         "\uC601\uC5C5\uD300 \uAC70\uB798\uCC98\uB294 \uC804\uAC74 \uBB34\uB8CC\uBC30\uC1A1. \uAC1C\uC778 \uAC70\uB798\uCC98\uB294 \uC81C\uD488\uAD70\uBCC4 \uCD5C\uC18C \uC218\uB7C9 \uAE30\uC900."],
        ["05", "\uBC30\uC1A1\uC9C0 \u2260 \uC5C5\uCCB4 \uC8FC\uC18C",
         "\uB300\uB9AC\uC810\u2192\uBCD1\uC6D0 \uC9C1\uC1A1 \uC2DC \uAC70\uB798\uBA85\uC138\uC11C \uBE44\uACE0\uB780\uC5D0 \uC2E4\uC81C \uBC30\uC1A1 \uBCD1\uC6D0\uBA85 \uC790\uB3D9 \uAE30\uC785."],
    ]
    create_table(doc, br_order_headers, br_order_rows, col_widths=[1.2, 3.8, 11.0])

    add_paragraph(doc, "", size=6, space_after=0)

    # ── 정산 규칙 ──
    add_paragraph(doc, "\uC815\uC0B0 \uADDC\uCE59", size=14, bold=True, color=PRIMARY,
                  space_before=12, space_after=6)

    br_settle_headers = ["BR", "\uADDC\uCE59\uBA85", "\uC0C1\uC138"]
    br_settle_rows = [
        ["06", "\uAC70\uB798\uCC98\uBCC4 \uB9C8\uAC10 \uAE30\uAC04",
         "\uB300\uB9AC\uC810: 1~25\uC77C(\uB610\uB294 \uB9D0\uC77C). \uBCD1\uC6D0: \uC804\uB2EC 26~\uB2F9\uC6D4 25. \uAC70\uB798\uCC98\uBCC4 \uAC1C\uBCC4 \uC124\uC815 \uAC00\uB2A5."],
        ["07", "\uACB0\uC81C \uBC29\uC2DD\uBCC4 \uCC98\uB9AC",
         "\uACC4\uC88C\uC774\uCCB4 \u2192 \uACC4\uC0B0\uC11C \uBC1C\uD589. \uCE74\uB4DC \uACB0\uC81C \u2192 \uBBF8\uBC1C\uD589. \uC0AC\uC6A9\uB7C9 \uACB0\uC81C \u2192 \uC601\uC5C5\uD300 \uC785\uB825 \uD6C4 \uC815\uC0B0."],
        ["08", "\uC0AC\uC6A9\uB7C9 \uACB0\uC81C \uBCD1\uC6D0",
         "\uC601\uC5C5\uD300 \uC0AC\uC6A9\uB7C9 \uC785\uB825 \uC804\uAE4C\uC9C0 \uD574\uB2F9 \uBCD1\uC6D0 \uB9C8\uAC10 \uC6D0\uC7A5\uC5D0 \uAE08\uC561 \uBBF8\uD45C\uC2DC. \uC785\uB825 \uD6C4 \uC790\uB3D9 \uBC18\uC601."],
        ["09", "\uAC70\uB798\uBA85\uC138\uC11C \uC591\uC2DD",
         "\uAC70\uB798\uCC98 \uC124\uC815\uC5D0 \uB530\uB77C \"\uAC70\uB798\uCC98 \uC591\uC2DD\" \uB610\uB294 \"RTBIO \uC591\uC2DD\". \uB3D9\uBD09 \uC5EC\uBD80\uB3C4 \uAC70\uB798\uCC98\uBCC4 \uC124\uC815."],
    ]
    create_table(doc, br_settle_headers, br_settle_rows, col_widths=[1.2, 3.8, 11.0])

    add_paragraph(doc, "", size=6, space_after=0)

    # ── 재고 규칙 ──
    add_paragraph(doc, "\uC7AC\uACE0 \uADDC\uCE59", size=14, bold=True, color=PRIMARY,
                  space_before=12, space_after=6)

    br_inv_headers = ["BR", "\uADDC\uCE59\uBA85", "\uC0C1\uC138"]
    br_inv_rows = [
        ["10", "\uCD9C\uACE0 \uC2DC \uCC28\uAC10",
         "\uCD9C\uACE0 \uC0C1\uD0DC \"\uCD9C\uACE0 \uC644\uB8CC\"\uB85C \uBCC0\uACBD\uB420 \uB54C \uC7AC\uACE0 \uCC28\uAC10. \uBC1C\uC8FC \uC811\uC218 \uC2DC\uC810\uC774 \uC544\uB2D8."],
        ["11", "\uC0D8\uD50C \uCD9C\uACE0",
         "\uCD9C\uACE0 \uC720\uD615 \"\uC0D8\uD50C\"\uC740 \uB9E4\uCD9C \uBBF8\uBC18\uC601 + \uC7AC\uACE0\uB9CC \uCC28\uAC10 + \uC6D0\uAC00 \uAE30\uB85D \uB0A8\uAE40."],
        ["12", "\uC138\uD2B8 \uC0C1\uD488 \uACC4\uC0B0",
         "\uBC1C\uC8FC \uC218\uB7C9 1 = \uC138\uD2B8 \uC218\uB7C9(\uC608: \uD3B8\uCE21 2\uAC1C). \uC7AC\uACE0 \uCC28\uAC10\uC740 \uC2E4\uC81C \uAC1C\uC218 \uAE30\uC900."],
    ]
    create_table(doc, br_inv_headers, br_inv_rows, col_widths=[1.2, 3.8, 11.0])

    add_paragraph(doc, "", size=6, space_after=0)

    # ══════════════════════════════════════════════════════════════════════
    # 6. 개발 일정 (8주)
    # ══════════════════════════════════════════════════════════════════════
    add_section_header(doc, "6", "\uAC1C\uBC1C \uC77C\uC815 (8\uC8FC)")

    sched_headers = ["\uAE30\uAC04", "\uB2E8\uACC4", "\uC8FC\uC694 \uB0B4\uC6A9", "\uC0B0\uCD9C\uBB3C"]
    sched_rows = [
        ["Week 1~3", "1\uCC28 \uAC1C\uBC1C",
         "\uACBD\uC601\uC9C0\uC6D0\uD300 \uC790\uB3D9\uD654, \uAC70\uB798\uCC98 \uBC1C\uC8FC \uD3EC\uD138, DB \uC778\uD504\uB77C \uAD6C\uCD95",
         "\uBC1C\uC8FC\uD3FC, \uAD00\uB9AC \uB300\uC2DC\uBCF4\uB4DC, DB \uC124\uACC4\uC11C"],
        ["Week 4~5", "2\uCC28 \uAC1C\uBC1C",
         "\uD488\uC9C8\uAD00\uB9AC\uD300 \uC7AC\uACE0, \uC785\uCD9C\uACE0 \uAD00\uB9AC, \uC791\uC5C5 \uC0C1\uD0DC \uAD00\uB9AC",
         "\uC7AC\uACE0 \uBAA8\uB4C8, \uCD9C\uACE0 \uAD00\uB9AC \uD654\uBA74"],
        ["Week 6~7", "3\uCC28 \uAC1C\uBC1C",
         "\uC601\uC5C5\uBD80 \uB300\uC2DC\uBCF4\uB4DC, \uB9E4\uCD9C/\uBBF8\uC218\uAE08 \uC870\uD68C, \uBCF4\uACE0\uC11C \uC790\uB3D9\uD654",
         "\uC601\uC5C5 \uB300\uC2DC\uBCF4\uB4DC, \uBCF4\uACE0\uC11C \uBAA8\uB4C8"],
        ["Week 8", "\uD1B5\uD569 \uD14C\uC2A4\uD2B8",
         "QA \uBC0F \uBC84\uADF8 \uC218\uC815, \uB370\uC774\uD130 \uC774\uAD00, \uBC30\uD3EC \uBC0F \uAD50\uC721",
         "\uD14C\uC2A4\uD2B8 \uACB0\uACFC\uC11C, \uC0AC\uC6A9\uC790 \uB9E4\uB274\uC5BC"],
    ]
    create_table(doc, sched_headers, sched_rows, col_widths=[2.5, 2.5, 6.0, 5.0])

    add_paragraph(doc, "", size=6, space_after=0)

    # ══════════════════════════════════════════════════════════════════════
    # 7. 가격 옵션
    # ══════════════════════════════════════════════════════════════════════
    add_section_header(doc, "7", "\uAC00\uACA9 \uC635\uC158")

    price_headers = ["\uD56D\uBAA9", "Option A (\uC77C\uC2DC\uBD88)", "Option B (\uAD6C\uB3C5\uD615)"]
    price_rows = [
        ["\uCD08\uAE30 \uBE44\uC6A9", "2,000\uB9CC\uC6D0", "1,000\uB9CC\uC6D0"],
        ["\uACB0\uC81C \uBC29\uC2DD", "\uCC29\uC218\uAE08 1,000 + \uC794\uAE08 1,000", "\uCC29\uC218\uAE08 500 + \uC794\uAE08 500"],
        ["\uC6D4 \uBE44\uC6A9 (AI \uD3EC\uD568)", "\uC6D4 25\uB9CC\uC6D0 (\uC5F0 300\uB9CC\uC6D0)", "\uC6D4 60\uB9CC\uC6D0 (\uC5F0 720\uB9CC\uC6D0)"],
        ["AI \uD06C\uB808\uB527", "\uC6D4 5,000 \uD3EC\uD568", "\uC6D4 5,000 \uD3EC\uD568"],
        ["AI \uCD08\uACFC \uC2DC", "WindyFlo Pro \uAE30\uC900 \uCD08\uACFC\uBD84 \uBCC4\uB3C4 \uCCAD\uAD6C", "WindyFlo Pro \uAE30\uC900 \uCD08\uACFC\uBD84 \uBCC4\uB3C4 \uCCAD\uAD6C"],
        ["\uC720\uC9C0\uBCF4\uC218", "\uAD6C\uB3C5 \uAE30\uAC04 \uC804\uCCB4 \uC81C\uACF5", "\uAD6C\uB3C5 \uAE30\uAC04 \uC804\uCCB4 \uC81C\uACF5"],
        ["\uC57D\uC815 \uAE30\uAC04", "\uC5C6\uC74C", "36\uAC1C\uC6D4"],
        ["\uAE30\uB2A5 \uC5C5\uB370\uC774\uD2B8", "\uBCC4\uB3C4 \uACAC\uC801", "\uBCC4\uB3C4 \uACAC\uC801"],
    ]
    create_table(doc, price_headers, price_rows, col_widths=[4.0, 6.0, 6.0])

    add_paragraph(doc, "", size=4, space_after=0)

    # AI note callout
    add_callout_box(doc,
        "AI \uAE30\uB2A5\uC740 WindyFlo \uD50C\uB7AB\uD3FC \uAE30\uBC18\uC73C\uB85C \uAD6C\uCD95\uD558\uBA70, "
        "WindyFlo Pro \uC694\uAE08\uC81C(\uC6D4 5,000) \uAE30\uC900\uC73C\uB85C \uD06C\uB808\uB527\uC774 \uC81C\uACF5\uB429\uB2C8\uB2E4. "
        "\uC77C\uBC18\uC801\uC778 \uC6D4 \uC5C5\uBB34\uB7C9(\uBC1C\uC8FC 600\uAC74 \uAE30\uC900)\uC740 \uD3EC\uD568 \uD06C\uB808\uB527 \uB0B4\uC5D0\uC11C \uCDA9\uBD84\uD788 \uCC98\uB9AC\uB429\uB2C8\uB2E4. "
        "\uD3EC\uD568 \uD06C\uB808\uB527 \uCD08\uACFC \uC2DC \uC2E4\uBE44 \uAE30\uC900\uC73C\uB85C \uBCC4\uB3C4 \uCCAD\uAD6C\uD569\uB2C8\uB2E4."
    )

    add_paragraph(doc, "", size=6, space_after=0)

    # ══════════════════════════════════════════════════════════════════════
    # 8. 계약 조건
    # ══════════════════════════════════════════════════════════════════════
    add_section_header(doc, "8", "\uACC4\uC57D \uC870\uAC74")

    contract_items = [
        None,  # placeholder — 개발 범위 is rendered separately below
        ("\uACB0\uC81C \uC870\uAC74",
         "\uCC29\uC218\uAE08\uC740 \uACC4\uC57D \uCCB4\uACB0 \uC2DC, \uC794\uAE08\uC740 \uCD5C\uC885 \uB0A9\uD488 \uC2DC \uC9C0\uAE09\uD569\uB2C8\uB2E4. "
         "\uC6D4 \uAD6C\uB3C5\uB8CC\uB294 \uC11C\uBE44\uC2A4 \uC624\uD508 \uC77C\uBD80\uD130 \uB9E4\uC6D4 \uC790\uB3D9 \uACB0\uC81C\uB429\uB2C8\uB2E4."),
        ("\uC9C0\uC801 \uC7AC\uC0B0\uAD8C",
         "\uAC1C\uBC1C \uC644\uB8CC \uBC0F \uC794\uAE08 \uC9C0\uAE09 \uC2DC \uC18C\uC2A4\uCF54\uB4DC \uBC0F \uC0B0\uCD9C\uBB3C\uC758 \uC18C\uC720\uAD8C\uC740 RTBIO\uC5D0 \uC774\uC804\uB429\uB2C8\uB2E4. "
         "\uB2E8, \uD504\uB808\uC784\uC6CC\uD06C \uBC0F \uC624\uD508\uC18C\uC2A4 \uB77C\uC774\uBE0C\uB7EC\uB9AC\uB294 \uAC01 \uB77C\uC774\uC120\uC2A4 \uC815\uCC45\uC744 \uB530\uB985\uB2C8\uB2E4."),
        ("\uC720\uC9C0\uBCF4\uC218",
         "\uAD6C\uB3C5 \uAE30\uAC04 \uB3D9\uC548 \uBC84\uADF8 \uC218\uC815 \uBC0F \uC11C\uBC84 \uC6B4\uC601\uC744 \uD3EC\uD568\uD569\uB2C8\uB2E4. "
         "\uC2E0\uADDC \uAE30\uB2A5 \uAC1C\uBC1C\uC740 \uBCC4\uB3C4 \uACAC\uC801\uC744 \uD1B5\uD574 \uD611\uC758\uD569\uB2C8\uB2E4."),
        ("\uBE44\uBC00\uC720\uC9C0",
         "\uC591\uCE21\uC740 \uD504\uB85C\uC81D\uD2B8 \uC9C4\uD589 \uC911 \uC54C\uAC8C \uB41C \uC0C1\uB300\uBC29\uC758 \uC601\uC5C5\uBE44\uBC00, \uAE30\uC220\uC815\uBCF4 \uB4F1\uC744 "
         "\uC81C3\uC790\uC5D0\uAC8C \uB204\uC124\uD558\uC9C0 \uC54A\uC2B5\uB2C8\uB2E4."),
        ("\uACC4\uC57D \uD574\uC9C0",
         "\uC77C\uBC29\uC801 \uACC4\uC57D \uD574\uC9C0 \uC2DC \uAE30 \uB0A9\uBD80 \uAE08\uC561\uC740 \uBC18\uD658\uB418\uC9C0 \uC54A\uC73C\uBA70, "
         "\uAC1C\uBC1C \uC644\uB8CC\uBD84\uC5D0 \uB300\uD55C \uC815\uC0B0\uC740 \uBCC4\uB3C4 \uD611\uC758\uD569\uB2C8\uB2E4."),
        ("\uBD84\uC7C1 \uD574\uACB0",
         "\uBCF8 \uACC4\uC57D\uACFC \uAD00\uB828\uD55C \uBD84\uC7C1\uC740 \uC6B0\uC120 \uD611\uC758\uB85C \uD574\uACB0\uD558\uBA70, "
         "\uD611\uC758 \uBD88\uBC1C \uC2DC \uAD00\uD560 \uBC95\uC6D0\uC744 \uD1B5\uD574 \uD574\uACB0\uD569\uB2C8\uB2E4."),
        ("\uAE30\uD0C0",
         "\uBCF8 \uACFC\uC5C5\uC9C0\uC2DC\uC11C\uC5D0 \uBA85\uC2DC\uB418\uC9C0 \uC54A\uC740 \uC0AC\uD56D\uC740 \uC591\uCE21 \uD611\uC758\uB97C \uD1B5\uD574 \uACB0\uC815\uD569\uB2C8\uB2E4."),
    ]

    # ── 1. 개발 범위 (expanded with change criteria) ──
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)
    run_num = p.add_run("  1. \uAC1C\uBC1C \uBC94\uC704  ")
    set_run_font(run_num, size=10, bold=True, color=PRIMARY)
    run_desc = p.add_run(
        "\uBCF8 \uACFC\uC5C5\uC9C0\uC2DC\uC11C\uC5D0 \uBA85\uC2DC\uB41C \uAE30\uB2A5\uC744 \uAE30\uBCF8 \uBC94\uC704\uB85C \uD558\uBA70, "
        "2\uCC28 \uBBF8\uD305 \uD6C4 \uC138\uBD80 \uD654\uBA74 \uAD6C\uC131 \uBC0F UX\uB97C \uCD5C\uC885 \uD655\uC815\uD569\uB2C8\uB2E4."
    )
    set_run_font(run_desc, size=10, bold=False, color=BLACK)

    # Change criteria table
    chg_headers = ["\uAD6C\uBD84", "\uAE30\uC900", "\uC608\uC2DC"]
    chg_rows = [
        ["\uBB34\uC0C1 \uBCC0\uACBD\n(\uC18C\uADDC\uBAA8)",
         "\uAC74\uB2F9 \uAC1C\uBC1C \uACF5\uC218 2\uC77C(16\uC2DC\uAC04) \uC774\uB0B4",
         "UI \uC218\uC815, \uC815\uB82C/\uD544\uD130 \uC870\uAC74 \uBCC0\uACBD,\n\uB9AC\uD3EC\uD2B8 \uC591\uC2DD \uBCC0\uACBD, \uD14D\uC2A4\uD2B8 \uC218\uC815"],
        ["\uC720\uC0C1 \uBCC0\uACBD\n(\uB300\uADDC\uBAA8)",
         "\uAC74\uB2F9 \uAC1C\uBC1C \uACF5\uC218 2\uC77C \uCD08\uACFC\n\u2192 \uBCC4\uB3C4 \uACAC\uC801\uC11C \uBC1C\uD589 \uD6C4 \uCC29\uC218",
         "\uC2E0\uADDC \uD654\uBA74/\uAE30\uB2A5 \uCD94\uAC00, DB \uAD6C\uC870 \uBCC0\uACBD,\n\uC678\uBD80 \uC2DC\uC2A4\uD15C \uC5F0\uB3D9 (\uCE74\uD1A1 \uC54C\uB9BC\uD1A1, \uD0DD\uBC30 API \uB4F1)"],
    ]
    create_table(doc, chg_headers, chg_rows, col_widths=[3.0, 5.5, 7.5])

    # Common rules as sub-paragraph
    p_common = doc.add_paragraph()
    p_common.paragraph_format.space_before = Pt(6)
    p_common.paragraph_format.space_after = Pt(4)
    # Indent
    pPr = p_common._element.get_or_add_pPr()
    ind = parse_xml(f'<w:ind {nsdecls("w")} w:left="200"/>')
    pPr.append(ind)
    items = [
        "\uBCC0\uACBD \uC694\uCCAD\uC740 \uC774\uBA54\uC77C \uB610\uB294 \uCE74\uCE74\uC624\uD1A1\uC73C\uB85C \uC811\uC218",
        "\uBB34\uC0C1/\uC720\uC0C1 \uD310\uB2E8\uC740 \uD558\uB9C8\uB2E4\uB7A9\uC2A4\uAC00 \uACF5\uC218 \uC0B0\uC815 \uD6C4 RTBIO\uC5D0 \uC0AC\uC804 \uACE0\uC9C0",
        "\uAD6C\uB3C5 \uAE30\uAC04 \uC911 \uBB34\uC0C1 \uBCC0\uACBD \uD69F\uC218 \uC81C\uD55C \uC5C6\uC74C (\uB2E8, \uB3D9\uC2DC \uC9C4\uD589 1\uAC74)",
    ]
    for j, item in enumerate(items):
        if j > 0:
            run_br = p_common.add_run("\n")
            set_run_font(run_br, size=9)
        run_bullet = p_common.add_run(f"\u2022 {item}")
        set_run_font(run_bullet, size=9, bold=False, color=RGBColor(0x33, 0x33, 0x33))

    # ── 2~8. Remaining contract items ──
    for i, item in enumerate(contract_items, 1):
        if item is None:
            continue  # skip placeholder (개발 범위 already rendered)
        label, desc = item
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(4)
        run_num = p.add_run(f"  {i}. {label}  ")
        set_run_font(run_num, size=10, bold=True, color=PRIMARY)
        run_desc = p.add_run(desc)
        set_run_font(run_desc, size=10, bold=False, color=BLACK)

    add_paragraph(doc, "", size=6, space_after=0)

    # ══════════════════════════════════════════════════════════════════════
    # 9. 납품물 목록
    # ══════════════════════════════════════════════════════════════════════
    add_section_header(doc, "9", "\uB0A9\uD488\uBB3C \uBAA9\uB85D")

    deliv_headers = ["No", "\uC0B0\uCD9C\uBB3C", "\uC124\uBA85", "\uB0A9\uD488 \uC2DC\uC810"]
    deliv_rows = [
        ["1", "\uC18C\uC2A4\uCF54\uB4DC", "\uC804\uCCB4 \uC2DC\uC2A4\uD15C \uC18C\uC2A4\uCF54\uB4DC (Git \uC800\uC7A5\uC18C)", "\uCD5C\uC885 \uB0A9\uD488 \uC2DC"],
        ["2", "DB \uC124\uACC4\uC11C (ERD)", "\uB370\uC774\uD130\uBCA0\uC774\uC2A4 \uAD6C\uC870 \uBB38\uC11C", "1\uCC28 \uB0A9\uD488 \uC2DC"],
        ["3", "API \uBB38\uC11C", "\uBC31\uC5D4\uB4DC API \uBA85\uC138", "\uCD5C\uC885 \uB0A9\uD488 \uC2DC"],
        ["4", "\uC0AC\uC6A9\uC790 \uB9E4\uB274\uC5BC", "\uAD00\uB9AC\uC790\uC6A9 / \uAC70\uB798\uCC98\uC6A9 \uAC01\uAC01", "\uCD5C\uC885 \uB0A9\uD488 \uC2DC"],
        ["5", "\uD14C\uC2A4\uD2B8 \uACB0\uACFC\uC11C", "\uAE30\uB2A5\uBCC4 \uD14C\uC2A4\uD2B8 \uD56D\uBAA9 \uBC0F \uACB0\uACFC", "\uCD5C\uC885 \uB0A9\uD488 \uC2DC"],
        ["6", "\uC6B4\uC601 \uAC00\uC774\uB4DC", "\uC11C\uBC84 \uC6B4\uC601, \uBC30\uD3EC, \uC7A5\uC560 \uB300\uC751 \uAC00\uC774\uB4DC", "\uCD5C\uC885 \uB0A9\uD488 \uC2DC"],
    ]
    create_table(doc, deliv_headers, deliv_rows, col_widths=[1.5, 3.5, 6.5, 4.5])

    add_paragraph(doc, "", size=6, space_after=0)

    # ══════════════════════════════════════════════════════════════════════
    # 10. 검수 기준 및 절차
    # ══════════════════════════════════════════════════════════════════════
    add_page_break(doc)
    add_section_header(doc, "10", "\uAC80\uC218 \uAE30\uC900 \uBC0F \uC808\uCC28")

    inspect_headers = ["\uB2E8\uACC4", "\uB0B4\uC6A9", "\uAE30\uC900", "\uB2F4\uB2F9"]
    inspect_rows = [
        ["1\uB2E8\uACC4\n\uAE30\uB2A5 \uAC80\uC218",
         "\uACFC\uC5C5\uC9C0\uC2DC\uC11C \uAE30\uB2A5 \uBAA9\uB85D(R01~R23) \uAE30\uBC18\n\uB3D9\uC791 \uD655\uC778 \uBC0F \uD14C\uC2A4\uD2B8",
         "\uC804\uCCB4 \uAE30\uB2A5 \uC815\uC0C1 \uC791\uB3D9",
         "\uD558\uB9C8\uB2E4\uB7A9\uC2A4"],
        ["2\uB2E8\uACC4\n\uB370\uC774\uD130 \uAC80\uC218",
         "\uAE30\uC874 \uB370\uC774\uD130 \uC774\uAD00 \uC815\uD655\uC131 \uD655\uC778\n(\uAC70\uB798\uCC98/\uC81C\uD488/\uD560\uC778\uC728 \uB4F1)",
         "\uC6D0\uBCF8 \uB370\uC774\uD130\uC640 100% \uC77C\uCE58",
         "RTBIO \uACBD\uC601\uC9C0\uC6D0\uD300"],
        ["3\uB2E8\uACC4\n\uC0AC\uC6A9\uC790 \uD14C\uC2A4\uD2B8",
         "\uC2E4\uC81C \uC5C5\uBB34 \uD658\uACBD\uC5D0\uC11C 2\uC8FC\uAC04 \uBCD1\uD589 \uC6B4\uC601\n(\uAE30\uC874 \uC5BC\uB9C8\uC5D0\uC694 + \uC2E0\uADDC \uC2DC\uC2A4\uD15C)",
         "\uC8FC\uC694 \uC5C5\uBB34 \uD750\uB984 \uC815\uC0C1 \uC218\uD589",
         "RTBIO \uC804 \uBD80\uC11C"],
        ["4\uB2E8\uACC4\n\uCD5C\uC885 \uC2B9\uC778",
         "\uAC80\uC218 \uACB0\uACFC \uD655\uC778 \uBC0F \uB0A9\uD488 \uC2B9\uC778",
         "RTBIO \uB300\uD45C \uC11C\uBA85",
         "RTBIO \uB300\uD45C"],
    ]
    create_table(doc, inspect_headers, inspect_rows, col_widths=[2.5, 5.5, 4.5, 3.5])

    add_paragraph(doc, "", size=4, space_after=0)

    add_callout_box(doc,
        "\uAC80\uC218 \uAE30\uAC04\uC740 \uCD5C\uC885 \uB0A9\uD488 \uD6C4 2\uC8FC\uC785\uB2C8\uB2E4. "
        "\uACBD\uBBF8\uD55C \uBC84\uADF8\uB294 \uAC80\uC218 \uC644\uB8CC \uD6C4 \uC720\uC9C0\uBCF4\uC218\uB85C \uCC98\uB9AC\uD558\uBA70, "
        "\uC911\uB300 \uACB0\uD568(\uB370\uC774\uD130 \uC190\uC2E4, \uC2DC\uC2A4\uD15C \uC7A5\uC560) \uBC1C\uACAC \uC2DC \uC218\uC815 \uC644\uB8CC \uD6C4 \uC7AC\uAC80\uC218\uB97C \uC9C4\uD589\uD569\uB2C8\uB2E4."
    )

    add_paragraph(doc, "", size=6, space_after=0)

    # ══════════════════════════════════════════════════════════════════════
    # 11. SLA 및 장애 대응
    # ══════════════════════════════════════════════════════════════════════
    add_section_header(doc, "11", "SLA \uBC0F \uC7A5\uC560 \uB300\uC751")

    sla_headers = ["\uAD6C\uBD84", "\uD56D\uBAA9", "\uAE30\uC900"]
    sla_rows = [
        ["\uC11C\uBE44\uC2A4 \uAC00\uC6A9\uB960", "\uC6D4\uAC04 \uAC00\uC6A9\uB960",
         "99.5% (\uC6D4 \uC57D 3.6\uC2DC\uAC04 \uC774\uB0B4 \uB2E4\uC6B4\uD0C0\uC784)"],
        ["\uC815\uAE30 \uC810\uAC80", "\uC11C\uBC84 \uC810\uAC80",
         "\uC0AC\uC804 \uACE0\uC9C0 \uD6C4 \uC9C4\uD589 (\uC5C5\uBB34 \uC678 \uC2DC\uAC04)"],
        ["\uC7A5\uC560 \uB4F1\uAE09", "\uAE34\uAE09 \u2014 \uC11C\uBE44\uC2A4 \uC804\uBA74 \uC911\uB2E8",
         "2\uC2DC\uAC04 \uB0B4 \uB300\uC751, 8\uC2DC\uAC04 \uB0B4 \uBCF5\uAD6C"],
        ["", "\uC911\uC694 \u2014 \uC8FC\uC694 \uAE30\uB2A5 \uC7A5\uC560",
         "4\uC2DC\uAC04 \uB0B4 \uB300\uC751, 24\uC2DC\uAC04 \uB0B4 \uBCF5\uAD6C"],
        ["", "\uC77C\uBC18 \u2014 \uACBD\uBBF8\uD55C \uC624\uB958",
         "\uC601\uC5C5\uC77C \uAE30\uC900 2\uC77C \uB0B4 \uB300\uC751"],
        ["\uC5F0\uB77D \uCC44\uB110", "\uC7A5\uC560 \uC2E0\uACE0",
         "\uC774\uBA54\uC77C + \uCE74\uCE74\uC624\uD1A1 (\uC6B4\uC601\uC2DC\uAC04: \uD3C9\uC77C 09~18\uC2DC)"],
        ["\uB370\uC774\uD130 \uBC31\uC5C5", "\uC790\uB3D9 \uBC31\uC5C5",
         "\uC77C 1\uD68C, \uCD5C\uADFC 30\uC77C \uBCF4\uAD00"],
        ["\uC7AC\uD574 \uBCF5\uAD6C", "DR \uACC4\uD68D",
         "\uBC31\uC5C5 \uAE30\uBC18 4\uC2DC\uAC04 \uB0B4 \uBCF5\uAD6C \uBAA9\uD45C"],
    ]
    create_table(doc, sla_headers, sla_rows, col_widths=[3.0, 5.0, 8.0])

    add_paragraph(doc, "", size=4, space_after=0)

    add_callout_box(doc,
        "SLA\uB294 \uAD6C\uB3C5 \uAE30\uAC04 \uC911 \uC801\uC6A9\uB418\uBA70, "
        "\uBD88\uAC00\uD56D\uB825(\uCC9C\uC7AC\uC9C0\uBCC0, AWS \uC7A5\uC560 \uB4F1)\uC740 \uAC00\uC6A9\uB960 \uC0B0\uC815\uC5D0\uC11C \uC81C\uC678\uB429\uB2C8\uB2E4. "
        "\uC6D4\uAC04 \uAC00\uC6A9\uB960\uC774 99.5% \uBBF8\uB9CC\uC77C \uACBD\uC6B0, \uD574\uB2F9 \uC6D4 \uAD6C\uB3C5\uB8CC\uC758 10%\uB97C \uD06C\uB808\uB527\uC73C\uB85C \uBCF4\uC0C1\uD569\uB2C8\uB2E4."
    )

    add_paragraph(doc, "", size=6, space_after=0)

    # ══════════════════════════════════════════════════════════════════════
    # 12. 보안 요구사항
    # ══════════════════════════════════════════════════════════════════════
    add_section_header(doc, "12", "\uBCF4\uC548 \uC694\uAD6C\uC0AC\uD56D")

    sec_headers = ["\uAD6C\uBD84", "\uC694\uAD6C\uC0AC\uD56D", "\uBE44\uACE0"]
    sec_rows = [
        ["\uB370\uC774\uD130 \uC804\uC1A1", "HTTPS(TLS 1.2+) \uC554\uD638\uD654 \uD1B5\uC2E0", "\uC804 \uAD6C\uAC04 \uC801\uC6A9"],
        ["\uC778\uC99D", "\uBE44\uBC00\uBC88\uD638 \uC554\uD638\uD654 \uC800\uC7A5 (bcrypt)", ""],
        ["\uC811\uADFC \uC81C\uC5B4", "\uBD80\uC11C\uBCC4/\uC5ED\uD560\uBCC4 \uAD8C\uD55C \uBD84\uB9AC",
         "\uC139\uC158 3 \uC0AC\uC6A9\uC790 \uC815\uC758 \uCC38\uC870"],
        ["\uAC70\uB798\uCC98 \uB370\uC774\uD130", "\uAC70\uB798\uCC98 \uAC04 \uB370\uC774\uD130 \uACA9\uB9AC\n(\uD0C0 \uAC70\uB798\uCC98 \uC815\uBCF4 \uC5F4\uB78C \uBD88\uAC00)", ""],
        ["\uAC10\uC0AC \uB85C\uADF8", "\uC8FC\uC694 \uC791\uC5C5 \uC774\uB825 \uAE30\uB85D\n(\uB85C\uADF8\uC778, \uBC1C\uC8FC, \uC0AD\uC81C, \uAD8C\uD55C \uBCC0\uACBD \uB4F1)",
         "\uCD5C\uC18C 1\uB144 \uBCF4\uAD00"],
        ["\uC11C\uBC84 \uBCF4\uC548", "AWS \uBCF4\uC548 \uADF8\uB8F9, \uBC29\uD654\uBCBD \uC124\uC815",
         "\uBD88\uD544\uC694\uD55C \uD3EC\uD2B8 \uCC28\uB2E8"],
        ["\uBC31\uC5C5", "\uC77C\uC77C \uC790\uB3D9 \uBC31\uC5C5 + 30\uC77C \uBCF4\uAD00",
         "SLA \uC139\uC158 \uCC38\uC870"],
    ]
    create_table(doc, sec_headers, sec_rows, col_widths=[3.0, 7.5, 5.5])

    add_paragraph(doc, "", size=6, space_after=0)

    # ══════════════════════════════════════════════════════════════════════
    # 13. RTBIO 전달 요청 사항
    # ══════════════════════════════════════════════════════════════════════
    add_section_header(doc, "13", "RTBIO \uC804\uB2EC \uC694\uCCAD \uC0AC\uD56D")

    req_headers = ["\uB2F4\uB2F9", "\uC804\uB2EC \uB0B4\uC6A9", "\uBE44\uACE0"]
    req_rows = [
        ["\uACBD\uC601\uC9C0\uC6D0\uD300", "\uC5BC\uB9C8\uC5D0\uC694 \uC0AC\uC6A9 \uC911\uC778 \uAE30\uB2A5 \uBAA9\uB85D", "\uC5B4\uB5A4 \uAE30\uB2A5\uC744 \uC5B8\uC81C \uC4F0\uB294\uC9C0"],
        ["\uACBD\uC601\uC9C0\uC6D0\uD300", "\uAE30\uC874 \uBC1C\uC8FC\uD3FC \uC5D1\uC140 \uD30C\uC77C", "\uB0B4\uC6A9 \uC5C6\uC774 \uC591\uC2DD\uB9CC"],
        ["\uACBD\uC601\uC9C0\uC6D0\uD300", "\uAC70\uB798\uCC98\uBCC4 \uD560\uC778\uC728/\uACB0\uC81C \uBC29\uC2DD/\uB9C8\uAC10 \uAE30\uAC04 \uC815\uB9AC", "\uC774\uBBF8 \uC815\uB9AC\uD574\uB454 \uAC83 \uC788\uC74C"],
        ["\uACBD\uC601\uC9C0\uC6D0\uD300", "\uC77C\uC77C \uBCF4\uACE0\uC11C/\uAE30\uAC04\uBCC4 \uB9E4\uCD9C \uBCF4\uACE0\uC11C \uC0D8\uD50C", "\uC5D1\uC140 \uD30C\uC77C"],
        ["\uACBD\uC601\uC9C0\uC6D0\uD300", "\uB9C8\uAC10 \uC6D0\uC7A5 \uC0D8\uD50C (\uB300\uB9AC\uC810\uC6A9/\uBCD1\uC6D0\uC6A9)", "\uC5D1\uC140 \uD30C\uC77C"],
        ["\uD488\uC9C8\uAD00\uB9AC\uD300", "\uCD94\uAC00 \uAE30\uB2A5 \uC694\uCCAD \uC0AC\uD56D", "\uC0DD\uAC01\uB098\uB294 \uB300\uB85C"],
        ["\uD488\uC9C8\uAD00\uB9AC\uD300", "\uC81C\uD488\uBCC4 \uBB34\uB8CC\uBC30\uC1A1 \uAE30\uC900 \uC218\uB7C9", "\uC81C\uD488\uAD70\uBCC4 \uC870\uAC74"],
    ]
    create_table(doc, req_headers, req_rows, col_widths=[3.0, 7.0, 6.0])

    add_paragraph(doc, "", size=6, space_after=0)

    # ══════════════════════════════════════════════════════════════════════
    # FOOTER
    # ══════════════════════════════════════════════════════════════════════
    add_paragraph(doc, "", size=10, space_after=0)
    add_paragraph(doc, "", size=10, space_after=0)

    # Divider
    p_div = doc.add_paragraph()
    p_div.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr = p_div._element.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="4" w:color="{LIGHT_GRAY}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)

    add_paragraph(doc, "\u00A9 2026 HAMADA LABS. All rights reserved.",
                  size=9, color=RGBColor(0x99, 0x99, 0x99),
                  alignment=WD_ALIGN_PARAGRAPH.CENTER,
                  space_before=8, space_after=0)

    # ── Save ──────────────────────────────────────────────────────────────
    doc.save(OUTPUT_PATH)
    return OUTPUT_PATH


# ── Entry point ───────────────────────────────────────────────────────────────
if __name__ == "__main__":
    path = build_document()
    print("Document generated successfully:")
    print(path)
