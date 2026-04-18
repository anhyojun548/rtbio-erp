# -*- coding: utf-8 -*-
"""
Patch RTBIO_과업지시서_v3.1.docx — 계약 조건 항목만 수정
빌드 스크립트 재생성 없이, 기존 파일을 열어서 특정 문단만 교체합니다.
Run: python -X utf8 patch_v31.py
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os, shutil

FILE = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                    "RTBIO_과업지시서_v3.1.docx")
BACKUP = FILE.replace(".docx", "_backup.docx")

FONT_NAME = "맑은 고딕"
PRIMARY = RGBColor(0x1B, 0x3A, 0x5C)
BLACK = RGBColor(0x00, 0x00, 0x00)


def set_run_font(run, size=10, bold=False, color=BLACK, font_name=FONT_NAME):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.name = font_name
    r = run._element
    rPr = r.find(qn('w:rPr'))
    if rPr is None:
        rPr = parse_xml(f'<w:rPr {nsdecls("w")}/>')
        r.insert(0, rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)


def clear_paragraph(p):
    """Remove all runs from a paragraph."""
    for run in p.runs:
        p._element.remove(run._element)
    # Also remove any remaining w:r elements
    for r_elem in p._element.findall(qn('w:r')):
        p._element.remove(r_elem)


def rewrite_contract_item(p, number, label, desc):
    """Rewrite a contract item paragraph: bold label + normal desc."""
    clear_paragraph(p)
    run_num = p.add_run(f"  {number}. {label}  ")
    set_run_font(run_num, size=10, bold=True, color=PRIMARY)
    run_desc = p.add_run(desc)
    set_run_font(run_desc, size=10, bold=False, color=BLACK)


def main():
    # Backup first
    shutil.copy2(FILE, BACKUP)
    print(f"백업 생성: {BACKUP}")

    doc = Document(FILE)

    # ── Find paragraph indices by matching text ──
    paragraphs = doc.paragraphs
    idx_map = {}
    for i, p in enumerate(paragraphs):
        t = p.text.strip()
        if t.startswith("2. 결제 조건"):
            idx_map['결제'] = i
        elif t.startswith("3. 지적 재산권"):
            idx_map['지재권'] = i
        elif t.startswith("6. 계약 해지"):
            idx_map['해지'] = i
        elif t.startswith("7. 분쟁 해결"):
            idx_map['분쟁'] = i
        elif t.startswith("8. 기타"):
            idx_map['기타'] = i

    print(f"수정 대상: {idx_map}")

    # ══════════════════════════════════════════════════════════════════════
    # 2. 결제 조건 — 잔금=검수완료 후, VAT 별도, 미납 시 처리
    # ══════════════════════════════════════════════════════════════════════
    p = paragraphs[idx_map['결제']]
    rewrite_contract_item(p, 2, "결제 조건",
        "모든 금액은 부가세(VAT) 별도입니다. "
        "착수금은 계약 체결 시, 잔금은 검수 완료일로부터 5영업일 이내에 지급합니다. "
        "검수 기간 내 서면 이의가 없을 경우 검수 완료로 간주합니다. "
        "월 구독료는 서비스 오픈일부터 매월 결제되며, "
        "30일 이상 미납 시 서면 통지 후 14일 유예 기간을 부여하고, "
        "미납이 지속될 경우 서비스를 중단할 수 있습니다."
    )
    print("  [완료] 2. 결제 조건")

    # ══════════════════════════════════════════════════════════════════════
    # 3. 지적 재산권 — 신규 개발분만 이전, 기존 자산 제외
    # ══════════════════════════════════════════════════════════════════════
    p = paragraphs[idx_map['지재권']]
    rewrite_contract_item(p, 3, "지적 재산권",
        "본 프로젝트를 위해 신규 개발된 산출물(소스코드, 화면 설계, 문서)의 "
        "저작재산권은 잔금 지급 완료 시 RTBIO에 이전됩니다. "
        "단, 하마다랩스의 기존 보유 프레임워크·범용 모듈·개발 도구 및 노하우는 "
        "이전 대상에서 제외되며, RTBIO 운영에 필요한 범위 내에서 비독점적 사용권을 부여합니다. "
        "오픈소스 및 제3자 라이브러리는 각 라이선스 정책을 따릅니다."
    )
    print("  [완료] 3. 지적 재산권")

    # ══════════════════════════════════════════════════════════════════════
    # 6. 계약 해지 — 귀책 분리, Option B 중도해지, 해지 후 처리
    # ══════════════════════════════════════════════════════════════════════
    p = paragraphs[idx_map['해지']]
    rewrite_contract_item(p, 6, "계약 해지",
        "해지를 원하는 측은 30일 전 서면 통지해야 합니다. "
        "① 상대방 귀책(중대한 계약 위반, 30일 이상 미납, 장기간 협조 불이행 등)으로 인한 해지 시, "
        "귀책 당사자가 상대방의 손해를 배상합니다. "
        "② Option B(구독형) 약정 기간 내 RTBIO 사유로 중도 해지 시, "
        "잔여 약정 기간 월 구독료의 50%를 위약금으로 정산합니다. "
        "③ 해지 후 처리: 하마다랩스는 해지일로부터 30일 이내에 "
        "데이터(DB 덤프)를 RTBIO에 반환하며, 반환 완료 후 30일 뒤 서버에서 삭제합니다. "
        "기 납부 금액은 귀책 사유 없는 한 반환되지 않습니다."
    )
    print("  [완료] 6. 계약 해지")

    # ══════════════════════════════════════════════════════════════════════
    # 7. 분쟁 해결 — 관할 법원 특정
    # ══════════════════════════════════════════════════════════════════════
    p = paragraphs[idx_map['분쟁']]
    rewrite_contract_item(p, 7, "분쟁 해결",
        "본 계약과 관련한 분쟁은 우선 협의로 해결하며, "
        "협의 불발 시 서울중앙지방법원을 제1심 전속 관할법원으로 합니다."
    )
    print("  [완료] 7. 분쟁 해결")

    # ══════════════════════════════════════════════════════════════════════
    # Save
    # ══════════════════════════════════════════════════════════════════════
    doc.save(FILE)
    print(f"\n저장 완료: {FILE}")
    print(f"파일 크기: {os.path.getsize(FILE)/1024:.1f} KB")


if __name__ == "__main__":
    main()
